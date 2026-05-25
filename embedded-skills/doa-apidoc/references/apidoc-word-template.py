# -*- coding: utf-8 -*-
"""
接口说明文档 Word 生成模板
=========================
使用方法：
  1. pip install python-docx
  2. 修改 build_docx() 中的接口内容数据
  3. python apidoc-word-template.py

此模板提供完整的基础设施（样式、组件函数），
只需替换 build_docx() 中的实际接口数据即可生成专业 Word 文档。
"""

import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


# ══════════════════════════════════════════════
# 常量
# ══════════════════════════════════════════════
BODY_FONT = "微软雅黑"
CODE_FONT = "Consolas"

C = {
    "pri": "1A365D",
    "sec": "2B6CB0",
    "acc": "3182CE",
    "lbg": "EBF4FF",
    "lgray": "F7FAFC",
    "border": "CBD5E0",
    "txt": "2D3748",
    "sub": "718096",
}


# ══════════════════════════════════════════════
# 基础样式函数
# ══════════════════════════════════════════════
def set_run_style(run, *, font_name=BODY_FONT, size=10.5, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_style(paragraph, *, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, line_spacing=1.25):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing


def clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]


def set_cell_shading(cell, fill):
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def set_cell_border(cell, color=C["border"], size=4):
    cell_properties = cell._tc.get_or_add_tcPr()
    borders = cell_properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        cell_properties.append(borders)

    for edge_name in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def append_markup(paragraph, text, *, font_name=BODY_FONT, size=10.5, color=None, bold=False):
    """支持 <b> 标记和 <br/> 换行的富文本"""
    normalized = text.replace("<br/>", "\n").replace("<br>", "\n")
    parts = re.split(r"(<b>.*?</b>)", normalized)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("<b>") and part.endswith("</b>")
        content = re.sub(r"</?b>", "", part)
        lines = content.split("\n")
        for index, line in enumerate(lines):
            run = paragraph.add_run(line)
            set_run_style(run, font_name=font_name, size=size, color=color, bold=bold or is_bold)
            if index < len(lines) - 1:
                run.add_break()


# ══════════════════════════════════════════════
# 高级组件函数
# ══════════════════════════════════════════════
def add_paragraph(doc, text, *, font_name=BODY_FONT, size=10.5, bold=False, color=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.3):
    paragraph = doc.add_paragraph()
    set_paragraph_style(paragraph, align=align, space_before=space_before, space_after=space_after, line_spacing=line_spacing)
    append_markup(paragraph, text, font_name=font_name, size=size, color=color, bold=bold)
    return paragraph


def add_section_bar(doc, text, color=None):
    """带色块的章节标题栏"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17)
    cell = table.cell(0, 0)
    set_cell_shading(cell, color or C["pri"])
    set_cell_border(cell, color or C["pri"], size=0)
    paragraph = clear_cell(cell)
    set_paragraph_style(paragraph, space_after=0, line_spacing=1.0)
    append_markup(paragraph, text, size=12, color="FFFFFF", bold=True)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_info_card(doc, rows):
    """信息卡片（封面元信息）"""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.8), Cm(13.2)]
    for label, value in rows:
        cells = table.add_row().cells
        for idx, width in enumerate(widths):
            cells[idx].width = width
            set_cell_shading(cells[idx], C["lbg"])
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_paragraph = clear_cell(cells[0])
        set_paragraph_style(label_paragraph, space_after=0, line_spacing=1.1)
        append_markup(label_paragraph, label, size=9.5, color=C["sub"], bold=True)
        value_paragraph = clear_cell(cells[1])
        set_paragraph_style(value_paragraph, space_after=0, line_spacing=1.1)
        append_markup(value_paragraph, value, size=9.5, color=C["txt"])


def add_divider(doc):
    """分隔线"""
    paragraph = doc.add_paragraph("_" * 90)
    set_paragraph_style(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.0)
    for run in paragraph.runs:
        set_run_style(run, size=6, color=C["border"])


def add_color_box(doc, text, *, fill=None, border=None):
    """调用场景高亮框"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill or C["lbg"])
    set_cell_border(cell, border or C["acc"], size=8)
    paragraph = clear_cell(cell)
    set_paragraph_style(paragraph, space_after=0, line_spacing=1.35)
    append_markup(paragraph, text, size=10.5, color=C["pri"], bold=True)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, col_widths_mm, *, header_fill=None):
    """标准数据表格（深蓝表头 + 交替行色）"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for index, width in enumerate(col_widths_mm):
        table.columns[index].width = Mm(width)

    header_cells = table.rows[0].cells
    for index, title in enumerate(headers):
        header_cells[index].width = Mm(col_widths_mm[index])
        set_cell_shading(header_cells[index], header_fill or C["pri"])
        set_cell_border(header_cells[index])
        paragraph = clear_cell(header_cells[index])
        set_paragraph_style(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.1)
        append_markup(paragraph, title, size=9, color="FFFFFF", bold=True)

    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            cells[col_index].width = Mm(col_widths_mm[col_index])
            if row_index % 2 == 0:
                set_cell_shading(cells[col_index], C["lgray"])
            set_cell_border(cells[col_index])
            paragraph = clear_cell(cells[col_index])
            alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_style(paragraph, align=alignment, space_after=0, line_spacing=1.15)
            append_markup(paragraph, str(value), size=9, color=C["txt"])
            cells[col_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_kv_table(doc, rows):
    """键值表格（左列浅蓝底粗体）"""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Mm(38), Mm(132)]
    for label, value in rows:
        cells = table.add_row().cells
        for idx, width in enumerate(widths):
            cells[idx].width = width
            set_cell_border(cells[idx], size=3)
        set_cell_shading(cells[0], C["lbg"])
        label_paragraph = clear_cell(cells[0])
        set_paragraph_style(label_paragraph, space_after=0, line_spacing=1.15)
        append_markup(label_paragraph, label, size=9, color=C["pri"], bold=True)
        value_paragraph = clear_cell(cells[1])
        set_paragraph_style(value_paragraph, space_after=0, line_spacing=1.15)
        append_markup(value_paragraph, value, size=9, color=C["txt"])


def add_code_block(doc, text):
    """代码块（浅灰底 + 等宽蓝色字体）"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.6)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8F9FA")
    set_cell_border(cell, size=4)
    paragraph = clear_cell(cell)
    set_paragraph_style(paragraph, space_after=0, line_spacing=1.1)
    run = paragraph.add_run(text.strip())
    set_run_style(run, font_name=CODE_FONT, size=8.5, color=C["sec"])


# ══════════════════════════════════════════════
# 文档配置
# ══════════════════════════════════════════════
def configure_document(header_left="项目名 · 接口说明文档",
                       header_right="CONFIDENTIAL 机密",
                       footer_left="公司名称",
                       footer_right="页码请在 Word 中查看"):
    """创建并配置 Word 文档（页面边距、字体、页眉页脚）"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = BODY_FONT
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal_style.font.size = Pt(10.5)

    # 页眉
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Cm(17))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_h = clear_cell(header_table.cell(0, 0))
    set_paragraph_style(left_h, space_after=0)
    append_markup(left_h, header_left, size=7.5, color=C["sub"])
    right_h = clear_cell(header_table.cell(0, 1))
    set_paragraph_style(right_h, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    append_markup(right_h, header_right, size=7.5, color=C["sub"])

    # 页脚
    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Cm(17))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_f = clear_cell(footer_table.cell(0, 0))
    set_paragraph_style(left_f, space_after=0)
    append_markup(left_f, footer_left, size=8, color=C["sub"])
    right_f = clear_cell(footer_table.cell(0, 1))
    set_paragraph_style(right_f, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    append_markup(right_f, footer_right, size=8, color=C["sub"])

    return doc


# ══════════════════════════════════════════════
# 构建 Word 文档（示例：替换为实际接口内容）
# ══════════════════════════════════════════════
def build_docx():
    # ── 配置区 ──
    filename = "接口说明文档.docx"
    doc = configure_document(
        header_left="项目名 · 接口说明文档",
        footer_left="公司名称",
    )

    # ══════ 封面 ══════
    add_paragraph(doc, "接 口 说 明 文 档", size=22, bold=True, color=C["pri"],
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_paragraph(doc, "项目名 — API Interface Specification", size=11,
                  color=C["sub"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_info_card(doc, [
        ("文档版本:", "V1.0"),
        ("编制方:", "XXXXXX"),
        ("接收方:", "XXXXXX"),
        ("编制日期:", "XXXX-XX-XX"),
        ("文档状态:", "初版发布"),
    ])
    add_paragraph(doc, "", space_after=2)
    add_divider(doc)

    # ══════ 一、接口总览 ══════
    add_section_bar(doc, "一、接口总览")
    add_paragraph(doc, "")
    add_paragraph(doc, "1.1 接口清单", size=11, bold=True, color=C["acc"], space_after=2)
    add_table(doc, ["#", "接口名称", "Method", "Path", "描述"],
              [["1", "示例接口", "GET", "/v1/example", "示例描述"]],
              [8, 26, 14, 64, 54])

    add_paragraph(doc, "1.2 调用链路", size=11, bold=True, color=C["acc"],
                  space_before=4, space_after=2)
    add_code_block(doc, """Client                          Server
  │                                │
  │  ① GET /v1/example             │
  │────────────────────────────────▶│
  │◀────────────────────────────────│
  │  {data: ...}                    │""")

    doc.add_page_break()

    # ══════ 二、公共约定 ══════
    add_section_bar(doc, "二、公共约定")
    add_paragraph(doc, "")
    add_paragraph(doc, "2.1 基础信息", size=11, bold=True, color=C["acc"], space_after=2)
    add_kv_table(doc, [
        ["协议", "HTTPS"],
        ["Base URL", "https://api.example.com/v1"],
        ["鉴权方式", "Bearer Token"],
        ["数据格式", "JSON（Content-Type: application/json; charset=utf-8）"],
    ])

    add_paragraph(doc, "2.2 公共响应结构", size=11, bold=True, color=C["acc"],
                  space_before=4, space_after=2)
    add_code_block(doc, """{
  "code":      200,
  "message":   "success",
  "data":      { ... },
  "timestamp": "2025-01-01T00:00:00+08:00"
}""")

    add_paragraph(doc, "2.3 错误码", size=11, bold=True, color=C["acc"],
                  space_before=4, space_after=2)
    add_table(doc, ["code", "含义", "说明"],
              [["200", "成功", "请求处理成功"],
               ["400", "参数错误", "请求参数缺失或格式非法"],
               ["401", "认证失败", "Token 无效或过期"],
               ["404", "资源不存在", "请求的资源不存在"],
               ["500", "服务内部错误", "服务端内部异常"]],
              [18, 30, 118])

    doc.add_page_break()

    # ══════ 三、接口详细定义 ══════
    add_section_bar(doc, "三、接口详细定义")
    add_paragraph(doc, "")
    add_paragraph(doc, "3.1  GET /v1/example — 示例接口", size=13, bold=True,
                  color=C["sec"], space_after=3)
    add_color_box(doc, "调用场景：示例调用场景说明。")
    add_paragraph(doc, "")

    add_paragraph(doc, "<b>入参（Query Parameter）</b>", size=11, bold=True,
                  color=C["acc"], space_after=2)
    add_table(doc, ["参数", "类型", "必填", "说明"],
              [["id", "string", "✅", "资源唯一标识"]],
              [30, 16, 12, 108])

    add_paragraph(doc, "<b>出参（Response data）</b>", size=11, bold=True,
                  color=C["acc"], space_before=4, space_after=2)
    add_table(doc, ["参数", "类型", "说明"],
              [["id", "string", "资源ID"],
               ["name", "string", "资源名称"]],
              [30, 18, 118])

    add_paragraph(doc, "<b>出参示例</b>", size=11, bold=True,
                  color=C["acc"], space_before=4, space_after=2)
    add_code_block(doc, """{
  "code": 200,
  "message": "success",
  "data": {
    "id": "RES-001",
    "name": "示例资源"
  }
}""")

    doc.add_page_break()

    # ══════ 四、调用时序总结 ══════
    add_section_bar(doc, "四、调用时序总结")
    add_paragraph(doc, "")
    add_table(doc, ["步骤", "调用方", "接口", "触发条件", "备注"],
              [["1", "Client → Server", "GET /v1/example", "用户触发", "返回资源详情"]],
              [12, 24, 52, 38, 40])

    # ══════ 保存 ══════
    doc.save(filename)
    print(f"✅ Word 文档已生成: {filename}")


if __name__ == "__main__":
    build_docx()
